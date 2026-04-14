"""
STEP12 멀티모델 AI 평가 스크립트 (한국어판) — 부동산 분석 보고서

AI 모델로 부동산 분석 보고서를 평가하고,
모든 항목 A 달성까지 반복 개선을 지원.

사용법:
    python 05_code/multi_model_evaluate.py <분석대상 ID>
    python 05_code/multi_model_evaluate.py seongsu-residential_KR
    python 05_code/multi_model_evaluate.py songpa-helio-city_KR

입력 파일 탐색 순서 (07_final_output_kr/):
    1. {대상ID}_designed.docx
    2. {대상ID}.docx
    3. (fallback) 06_middle_output/{대상ID}/STEP11_보고서_draft.md

환경변수 (.env 파일):
    ANTHROPIC_API_KEY  - Anthropic API 키 (Claude)
    OPENAI_API_KEY     - OpenAI API 키 (GPT)
    GEMINI_API_KEY     - Google AI (Gemini) API 키

의존성:
    pip install litellm anthropic python-dotenv python-docx
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from dotenv import load_dotenv
import litellm

# .env 로드 (프로젝트 루트)
ROOTDIR = Path(__file__).resolve().parent.parent
load_dotenv(ROOTDIR / ".env")

# API 키 trailing whitespace 제거 (개행문자 포함)
for _key in ("ANTHROPIC_API_KEY", "OPENAI_API_KEY", "GEMINI_API_KEY"):
    if os.environ.get(_key):
        os.environ[_key] = os.environ[_key].strip()

MODELS = {
    "Gemini-3-Flash": "gemini/gemini-3-flash-preview",
}

EVALUATION_PROMPT = """\
# 의뢰 내용
당신은 부동산 투자·개발 분야의 숙련된 애널리스트입니다.
아래 「부동산 분석 보고서」를 5개 관점으로 평가하고, 구체적 피드백을 제공해 주세요.

# 평가기준
아래 5개 관점으로 A(우수)·B(표준)·C(개선필요) 3단계 평가하고, 구체적 피드백을 제공해 주세요.

| 평가관점 | 평가항목 | 고평가 기준 |
|---|---|---|
| 1. 데이터 신뢰성 | 수치의 출처와 기준 시점 | 모든 주요 수치에 출처(국토부, KB, R-ONE 등)와 기준 시점이 명시됐는가 |
| 2. 시장 해석 | 가격·수급 분석의 논리 연결 | 가격 흐름과 수요·공급 변화가 인과관계로 설명됐는가 |
| 3. 입지 분석 | 입지와 수요의 연결 | 교통·학군·업무 접근성이 실수요층과 구체적으로 연결됐는가 |
| 4. 리스크 반영 | 정책·공급·금리 리스크 | 대출 규제·입주 물량·금리 변화 등 하방 리스크가 구조화됐는가 |
| 5. 결론 실효성 | 의사결정에의 실질적 도움 | 매수/보유/매도/개발/보류 등 행동 방향이 근거와 함께 명확히 제시됐는가 |

# 출력 형식
## 1. 종합평가
(A/B/C 랭크와 그 이유를 간결히)

## 2. 항목별 평가
- **데이터 신뢰성**: [랭크] 코멘트...
- **시장 해석**: [랭크] 코멘트...
- **입지 분석**: [랭크] 코멘트...
- **리스크 반영**: [랭크] 코멘트...
- **결론 실효성**: [랭크] 코멘트...

## 3. 강점 (Strengths)
-
-

## 4. 개선 조언 (Areas for Improvement)
-
-

# 부동산 분석 보고서
(아래에 보고서 본문을 첨부합니다)
"""

CRITERIA = ["데이터 신뢰성", "시장 해석", "입지 분석", "리스크 반영", "결론 실효성"]


def extract_bare_id(target_id: str) -> str:
    """분석대상 ID에서 '_KR' 접미사 제거 (예: 'seongsu-residential_KR' → 'seongsu-residential')"""
    return re.sub(r'_KR$', '', target_id)


def extract_docx_text(docx_path: Path) -> str:
    """DOCX 파일에서 텍스트 추출 (단락 + 표)"""
    doc = Document(str(docx_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    parts = paragraphs
    if table_lines:
        parts += [""] + table_lines

    return "\n".join(parts)


def find_report_text(target_id: str) -> tuple[str, str]:
    """
    보고서 텍스트와 소스 경로를 반환.
    탐색 순서:
      1. 07_final_output_kr/{대상ID}_designed.docx
      2. 07_final_output_kr/{대상ID}.docx
      3. 06_middle_output/{대상ID}/STEP11_보고서_draft.md (fallback)
    """
    bare = extract_bare_id(target_id)
    docx_candidates = [
        ROOTDIR / "07_final_output_kr" / f"{bare}_designed.docx",
        ROOTDIR / "07_final_output_kr" / f"{bare}.docx",
        ROOTDIR / "07_final_output_kr" / f"{target_id}_designed.docx",
        ROOTDIR / "07_final_output_kr" / f"{target_id}.docx",
    ]
    for path in docx_candidates:
        if path.exists():
            text = extract_docx_text(path)
            return text, str(path)

    # fallback: MD draft
    md_candidates = [
        ROOTDIR / "06_middle_output" / target_id / "STEP10_보고서_draft.md",
        ROOTDIR / "06_middle_output" / target_id / "STEP11_output.md",
    ]
    for path in md_candidates:
        if path.exists():
            return path.read_text(encoding="utf-8"), str(path)

    print(f"ERROR: 보고서를 찾을 수 없음: {target_id}")
    print("검색 경로:")
    for p in docx_candidates + md_candidates:
        print(f"  {p}")
    sys.exit(1)


def parse_grades(response_text: str) -> dict:
    """AI 응답에서 항목별 등급(A/B/C) 추출"""
    grades = {}
    for criterion in CRITERIA:
        escaped = re.escape(criterion)
        pattern = rf"{escaped}[^A-C]*[\[【]?([ABC])[\]】]?"
        match = re.search(pattern, response_text)
        grades[criterion] = match.group(1) if match else "?"
    return grades


def evaluate_with_model(model_id: str, prompt: str) -> str:
    """LiteLLM 경유 모델 호출"""
    response = litellm.completion(
        model=model_id,
        messages=[{"role": "user", "content": prompt}],
        timeout=600,
    )
    return response.choices[0].message.content


def run_evaluation(target_id: str) -> dict:
    """1회 반복 실행"""
    report_text, source_path = find_report_text(target_id)
    print(f"보고서: {source_path}")

    full_prompt = EVALUATION_PROMPT + "\n" + report_text

    results = {}
    for model_name, model_id in MODELS.items():
        print(f"  {model_name} ({model_id}) 평가 중...")
        try:
            response_text = evaluate_with_model(model_id, full_prompt)
            grades = parse_grades(response_text)
            results[model_name] = {
                "grades": grades,
                "full_response": response_text,
            }
            print(f"    결과: {grades}")
        except Exception as e:
            print(f"    ERROR: {e}")
            results[model_name] = {
                "grades": {c: "?" for c in CRITERIA},
                "full_response": f"Error: {e}",
            }
    return results


def check_all_a(results: dict) -> bool:
    for data in results.values():
        for grade in data["grades"].values():
            if grade != "A":
                return False
    return True


def format_summary_table(results: dict) -> str:
    model_names = list(MODELS.keys())
    header = "| 평가항목 | " + " | ".join(model_names) + " |"
    separator = "|----------|" + "|".join(["------" for _ in model_names]) + "|"
    lines = [header, separator]
    for criterion in CRITERIA:
        row = [criterion]
        for model_name in model_names:
            grade = results.get(model_name, {}).get("grades", {}).get(criterion, "?")
            row.append(grade)
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def save_results(target_id: str, iteration: int, results: dict, all_a: bool):
    output_dir = ROOTDIR / "06_middle_output" / target_id
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "STEP12_output.md"

    today = datetime.now().strftime("%Y-%m-%d")
    summary_table = format_summary_table(results)

    content = ""
    if iteration == 1:
        content += "\n\n## 멀티모델 AI 평가\n"

    content += f"\n### Iteration {iteration} ({today})\n"
    content += f"\n#### 평가결과 요약\n{summary_table}\n"

    # B/C 항목 (개선 대상)
    bc_items = []
    for model_name, data in results.items():
        for criterion, grade in data["grades"].items():
            if grade in ("B", "C"):
                bc_items.append(f"- [{model_name} / {criterion}: {grade}]")
    if bc_items:
        content += "\n#### 개선 필요 항목\n" + "\n".join(bc_items) + "\n"

    # 상세 응답
    content += "\n#### 각 모델 상세 응답\n"
    for model_name in MODELS:
        data = results.get(model_name, {})
        content += f"\n<details><summary>{model_name}</summary>\n\n"
        content += data.get("full_response", "(no response)") + "\n"
        content += "\n</details>\n"

    if all_a:
        content += "\n### 최종 결과\n모든 모델·모든 항목 A 달성: ✅\n"

    with open(output_path, "a", encoding="utf-8") as f:
        f.write(content)

    print(f"결과 저장: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="STEP12: 멀티모델 AI 평가 — 부동산 분석 보고서")
    parser.add_argument("target_id", help="분석대상 ID (예: seongsu-residential_KR)")
    parser.add_argument("--max-iterations", type=int, default=5, help="최대 반복 횟수 (기본 5)")
    args = parser.parse_args()

    target_id = args.target_id

    print(f"=== 멀티모델 AI 평가: {target_id} ===\n")

    for iteration in range(1, args.max_iterations + 1):
        print(f"\n--- Iteration {iteration} ---")
        results = run_evaluation(target_id)
        all_a = check_all_a(results)
        save_results(target_id, iteration, results, all_a)

        print(f"\n{format_summary_table(results)}")

        if all_a:
            print(f"\n[OK] 모든 모델·모든 항목 A 달성! (Iteration {iteration})")
            break
        else:
            print("\n[NG] A 미달 항목 존재. 보고서 수정 후 재실행 권장.")
            if iteration < args.max_iterations:
                try:
                    input("   [Enter] 다음 반복 / [Ctrl+C] 중단 > ")
                except (KeyboardInterrupt, EOFError):
                    print("\n중단됨.")
                    break
    else:
        print(f"\n[NG] {args.max_iterations}회 반복에도 전원 A 미달성.")


if __name__ == "__main__":
    main()
