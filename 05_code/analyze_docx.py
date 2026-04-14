"""
DOCX 제안서 구조 분석 스크립트

사용 예:
    python 05_code/analyze_docx.py 07_final_output_kr/035890_designed.docx
"""
import os
import sys

from docx import Document


def analyze_document(doc_path):
    """문서 구조를 분석해 요약 출력"""
    doc = Document(doc_path)

    print(f"=== 분석: {os.path.basename(doc_path)} ===\n")

    # 단락 정보
    print(f"총 단락수: {len(doc.paragraphs)}")
    print(f"표 수: {len(doc.tables)}")

    # 스타일 사용 현황
    style_usage = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"
        style_usage[style_name] = style_usage.get(style_name, 0) + 1

    print("\n=== 사용 스타일 ===")
    for style, count in sorted(style_usage.items(), key=lambda x: -x[1]):
        print(f"  {style}: {count}회")

    # 단락 샘플 (처음 30개)
    print("\n=== 단락 샘플 (처음 30개) ===")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "None"
            display_text = text[:50] + "..." if len(text) > 50 else text
            print(f"{i:3d}: [{style_name:15s}] {display_text}")

    # 표 정보
    if doc.tables:
        print("\n=== 표 정보 ===")
        for i, table in enumerate(doc.tables[:3]):
            print(f"표{i+1}: {len(table.rows)}행 x {len(table.columns)}열")
            if table.rows:
                first_row = [cell.text.strip()[:20] for cell in table.rows[0].cells]
                print(f"  헤더: {first_row}")

    # 섹션 정보
    print("\n=== 섹션 정보 ===")
    for i, section in enumerate(doc.sections):
        print(f"섹션{i+1}:")
        print(f"  페이지 크기: {section.page_width.inches:.2f} x {section.page_height.inches:.2f} 인치")
        print(
            f"  여백(상·하·좌·우): {section.top_margin.inches:.2f}, {section.bottom_margin.inches:.2f}, "
            f"{section.left_margin.inches:.2f}, {section.right_margin.inches:.2f} 인치"
        )

        # 머리글/바닥글
        header = section.header
        footer = section.footer
        header_text = "".join([p.text for p in header.paragraphs]).strip() if header.paragraphs else ""
        footer_text = "".join([p.text for p in footer.paragraphs]).strip() if footer.paragraphs else ""
        print(f"  머리글: {'있음' if header_text else '없음'}")
        print(f"  바닥글: {'있음' if footer_text else '없음'}")


if __name__ == "__main__":
    doc_path = sys.argv[1] if len(sys.argv) > 1 else "07_final_output_kr/035890_designed.docx"

    if os.path.exists(doc_path):
        analyze_document(doc_path)
    else:
        print(f"파일을 찾을 수 없습니다: {doc_path}")
