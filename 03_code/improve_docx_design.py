"""
제안서(.docx)의 디자인을 개선하는 스크립트

기능:
1. 제목 스타일 개선(색상, 폰트, 밑줄)
2. 표 디자인 개선(헤더 행 배경색, 테두리선)
3. 헤더/푸터 추가
4. 표지 디자인 개선
5. 페이지 번호 추가

사용법:
  python improve_docx_design.py [입력파일] [출력파일]
  python improve_docx_design.py --all  # 전체 파일을 일괄 처리
"""

import os
import sys
import re
import shutil
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

# ========================================
# 디자인 설정(커스터마이즈 가능)
# ========================================

# 컬러 테마(네이비 계열의 비즈니스 컬러)
COLORS = {
    'primary': RGBColor(0, 51, 102),       # 진한 네이비(제목, 강조)
    'secondary': RGBColor(51, 102, 153),   # 블루(서브 제목)
    'accent': RGBColor(0, 102, 153),       # 틸(강조)
    'table_header': RGBColor(0, 51, 102),  # 표 헤더 배경
    'table_header_text': RGBColor(255, 255, 255),  # 표 헤더 문자
    'table_alt_row': RGBColor(240, 245, 250),  # 표의 교차 행 배경
    'border': RGBColor(180, 180, 180),     # 테두리선
}

# 폰트 설정
FONTS = {
    'heading': '맑은고딕',
    'body': '맑은고딕',
    'title': '맑은고딕',
}

ROOTDIR = Path(__file__).resolve().parent.parent
KNOWN_AGENT_TAGS = {'claude', 'codex'}

# ========================================
# 헬퍼 함수
# ========================================

def set_cell_shading(cell, color):
    """셀의 배경색을 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, border_color='000000', border_size='4'):
    """셀의 테두리선을 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 기존의 테두리선 설정을 삭제
    for border in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(border)

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_page_number(run):
    """페이지 번호 필드를 추가"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_total_pages(run):
    """총 페이지 수 필드를 추가"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def extract_target_id(filepath):
    """입력 파일 경로에서 분석대상 ID를 추출"""
    path = Path(filepath)
    if path.parent.name:
        return path.parent.name

    filename = path.stem
    for suffix in ('_designed', '_draft'):
        if filename.endswith(suffix):
            filename = filename[:-len(suffix)]
    return filename


def extract_agent_tag_from_filename(filepath):
    """파일명에서 report_*_<agent>.docx 형태의 에이전트 태그를 추출"""
    stem = Path(filepath).stem
    match = re.search(r'_(claude|codex)$', stem)
    return match.group(1) if match else ""


def build_header_label(target_id):
    """헤더에 표시할 분석대상 식별자를 생성"""
    target_id = (target_id or "").strip()
    return target_id or "unknown-target"


def build_final_output_path(target_id, designed_path):
    """최종본을 05_output 아래에 배치할 경로를 생성"""
    target_id = (target_id or "").strip()
    final_dir = ROOTDIR / "05_output"
    final_dir.mkdir(parents=True, exist_ok=True)

    if target_id and target_id != "unknown-target":
        return final_dir / f"{target_id}_designed.docx"
    return final_dir / Path(designed_path).name

def extract_company_name(doc):
    """문서에서 표지의 대상명을 추출"""
    for para in doc.paragraphs[:20]:
        text = para.text.strip()
        if text.startswith('분석 대상:'):
            return text.split(':', 1)[1].strip()
        if text and text not in {'부동산 분석 보고서'} and '작성자:' not in text:
            return text
    return ""

# ========================================
# 디자인 개선 함수
# ========================================

def improve_heading_styles(doc):
    """제목 스타일을 개선"""
    print("  제목 스타일 개선 중...")

    # Heading 1 스타일(장 제목)
    try:
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = FONTS['heading']
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.font.color.rgb = COLORS['primary']
        h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['heading'])

        # 단락 포맷
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(6)
        h1_style.paragraph_format.keep_with_next = True

        # 밑줄(테두리) 추가
        pPr = h1_style._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt
        bottom.set(qn('w:color'), '003366')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception as e:
        print(f"    Heading 1 스타일 설정 오류: {e}")

    # Heading 2 스타일(섹션)
    try:
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = FONTS['heading']
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.font.color.rgb = COLORS['secondary']
        h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['heading'])

        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(4)
        h2_style.paragraph_format.keep_with_next = True
    except Exception as e:
        print(f"    Heading 2 스타일 설정 오류: {e}")

    # Normal 스타일
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = FONTS['body']
        normal_style.font.size = Pt(10)
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(6)
    except Exception as e:
        print(f"    Normal 스타일 설정 오류: {e}")

def improve_tables(doc):
    """표의 디자인을 개선"""
    print("  표 디자인 개선 중...")

    for table_idx, table in enumerate(doc.tables):
        try:
            # 표의 너비를 설정(페이지 너비에 맞춤)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 헤더 행(첫 번째 행)의 스타일
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # 배경색
                    set_cell_shading(cell, (COLORS['table_header'][0], COLORS['table_header'][1], COLORS['table_header'][2]))
                    # 테두리선
                    set_cell_border(cell, '003366', '8')
                    # 텍스트 스타일
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = COLORS['table_header_text']
                            run.font.size = Pt(10)

            # 데이터 행의 스타일
            for row_idx, row in enumerate(table.rows[1:], 1):
                for cell in row.cells:
                    # 교차 행의 배경색
                    if row_idx % 2 == 0:
                        set_cell_shading(cell, (COLORS['table_alt_row'][0], COLORS['table_alt_row'][1], COLORS['table_alt_row'][2]))
                    # 테두리선
                    set_cell_border(cell, 'B4B4B4', '4')
                    # 텍스트 스타일
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)

        except Exception as e:
            print(f"    표{table_idx+1}의 스타일 설정 오류: {e}")

def add_header_footer(doc, target_id, company_name=""):
    """헤더와 푸터를 추가"""
    print("  머리글/바닥글 추가 중...")

    for section_idx, section in enumerate(doc.sections):
        # 표지 섹션(첫 번째)은 건너뜀
        if section_idx == 0:
            # 첫 번째 섹션은 다른 헤더/푸터를 사용
            section.different_first_page_header_footer = True

        # 헤더
        header = section.header
        header.is_linked_to_previous = False

        # 헤더의 내용을 초기화
        for para in header.paragraphs:
            para.clear()

        # 헤더 텍스트를 추가
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        header_text = f"부동산 분석 보고서 | 분석대상: {target_id}"
        run = header_para.add_run(header_text)
        run.font.name = FONTS['heading']
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS['secondary']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['heading'])

        # 헤더 아래의 테두리선
        pPr = header_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '336699')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 푸터
        footer = section.footer
        footer.is_linked_to_previous = False

        # 푸터의 내용을 초기화
        for para in footer.paragraphs:
            para.clear()

        # 푸터 텍스트를 추가(페이지 번호)
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 페이지 번호: X / Y 형식
        run1 = footer_para.add_run()
        run1.font.name = FONTS['body']
        run1.font.size = Pt(9)
        run1.font.color.rgb = COLORS['secondary']
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])
        add_page_number(run1)

        run2 = footer_para.add_run(" / ")
        run2.font.name = FONTS['body']
        run2.font.size = Pt(9)
        run2.font.color.rgb = COLORS['secondary']
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

        run3 = footer_para.add_run()
        run3.font.name = FONTS['body']
        run3.font.size = Pt(9)
        run3.font.color.rgb = COLORS['secondary']
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])
        add_total_pages(run3)

def _add_left_border(paragraph, color_hex='003366', width_pt=4, space_pt=10):
    """단락의 왼쪽 끝에 세로 강조선을 추가"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(width_pt * 8))  # 1/8 pt 단위
    left.set(qn('w:space'), str(space_pt))
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def improve_cover_page(doc, target_id):
    """표지 디자인을 개선(안C: 좌측 정렬 비즈니스 리포트 스타일)"""
    print("  표지 디자인 개선 중...")

    # 표지 단락을 식별(제1장 전까지)
    cover_paragraphs = []
    for i, para in enumerate(doc.paragraphs[:25]):
        text = para.text.strip()
        if '제1장' in text or 'Executive Summary' in text or '에그제큐티브' in text:
            break
        cover_paragraphs.append((i, para))

    # 타이틀 블록(세로선을 붙이는 대상)의 인덱스를 기록
    title_indices = []

    for i, para in cover_paragraphs:
        text = para.text.strip()
        if not text:
            continue

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 메인 타이틀
        if '부동산 분석 보고서' in text:
            for run in para.runs:
                run.font.name = FONTS['title']
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['title'])
            para.paragraph_format.space_after = Pt(4)
            title_indices.append(i)

        # 분석 대상명
        elif text.startswith('분석 대상:'):
            for run in para.runs:
                run.font.name = FONTS['title']
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['title'])
            para.paragraph_format.space_after = Pt(2)

        # 분석대상 ID
        elif target_id and target_id in text:
            for run in para.runs:
                run.font.name = FONTS['body']
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS['secondary']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

        # 작성자·날짜(하단 메타 정보)
        elif '작성자' in text or ('년' in text and len(text) < 20):
            for run in para.runs:
                run.font.name = FONTS['body']
                run.font.size = Pt(10)
                run.font.color.rgb = COLORS['secondary']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

        # 서브 타이틀(제안서 제목 등, 위에 해당하지 않는 긴 텍스트)
        else:
            for run in para.runs:
                run.font.name = FONTS['title']
                run.font.size = Pt(13)
                run.font.italic = True
                run.font.color.rgb = COLORS['secondary']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['title'])
            para.paragraph_format.space_before = Pt(6)
            title_indices.append(i)

    # 타이틀 블록에 왼쪽 세로선을 추가
    for idx in title_indices:
        _add_left_border(doc.paragraphs[idx])

def improve_paragraphs(doc):
    """본문 단락의 디자인을 개선"""
    print("  본문 단락 개선 중...")

    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ""

        # 이미 Heading 스타일인 경우 건너뜀(스타일 설정에서 처리 완료)
        if style_name.startswith('Heading'):
            continue

        # 항목 나열의 강조(■, ●, ◆로 시작하는 행)
        if text.startswith('■') or text.startswith('●') or text.startswith('◆'):
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = COLORS['primary']
                run.font.size = Pt(10)

        # 하이픈이나 불릿의 항목 나열
        elif text.startswith('-'):
            para.paragraph_format.left_indent = Cm(0.5)

        # 일반 본문
        else:
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(10)
                run.font.name = FONTS['body']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONTS['body'])

def process_document(input_path, output_path=None):
    """문서 전체를 처리"""
    print(f"\n처리 중: {input_path}")

    if output_path is None:
        input_file = Path(input_path)
        agent_tag = extract_agent_tag_from_filename(input_file.name)
        if input_file.name == "report_draft.docx":
            output_path = str(input_file.with_name("report_designed.docx"))
        elif agent_tag and input_file.name == f"report_draft_{agent_tag}.docx":
            output_path = str(input_file.with_name(f"report_designed_{agent_tag}.docx"))
        else:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_designed{ext}"

    # 문서를 읽어들임
    doc = Document(input_path)

    # 분석대상 ID/이름 추출
    target_id = build_header_label(extract_target_id(input_path))
    company_name = extract_company_name(doc)
    print(f"  분석대상 ID: {target_id}")
    print(f"  분석대상명: {company_name if company_name else '(검출 실패)'}")

    # 디자인 개선을 적용
    improve_heading_styles(doc)
    improve_tables(doc)
    improve_cover_page(doc, target_id)
    improve_paragraphs(doc)
    add_header_footer(doc, target_id, company_name)

    # 저장 중 다른 프로세스가 부분 저장본을 읽지 않도록 임시 파일에 쓴 뒤 교체
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(
        prefix=f".{output_path_obj.stem}_",
        suffix=output_path_obj.suffix,
        dir=str(output_path_obj.parent),
    )
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, output_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    print(f"  저장 완료: {output_path}")

    final_output_path = build_final_output_path(target_id, output_path)
    shutil.copy2(output_path, final_output_path)
    print(f"  최종본 배치 완료: {final_output_path}")

    return str(final_output_path)

def process_all_documents():
    """04_workspace 하위의 draft DOCX를 일괄 처리"""
    workspace_dir = ROOTDIR / "04_workspace"
    docx_files = sorted(workspace_dir.glob("*/report_draft*.docx"))

    print(f"처리 대상: {len(docx_files)}파일")

    for input_path in docx_files:
        agent_tag = extract_agent_tag_from_filename(input_path.name)
        if agent_tag:
            output_path = input_path.with_name(f"report_designed_{agent_tag}.docx")
        else:
            output_path = input_path.with_name("report_designed.docx")
        try:
            process_document(str(input_path), str(output_path))
        except Exception as e:
            print(f"  오류 ({input_path}): {e}")

    print(f"\n완료! 작업본 경로: {workspace_dir}")
    print(f"최종본 경로: {ROOTDIR / '05_output'}")

# ========================================
# 메인
# ========================================

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법:")
        print("  python improve_docx_design.py [입력파일.docx] [출력파일.docx]")
        print("  python improve_docx_design.py --all  # 전체 파일을 일괄 처리")
        sys.exit(1)

    if sys.argv[1] == "--all":
        process_all_documents()
    else:
        input_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        process_document(input_path, output_path)
