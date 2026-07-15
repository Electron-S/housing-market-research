"""mermaid로 생성한 이미지를 DOCX 부동산 분석 보고서에 삽입한다.

사용 예:
    python 03_code/insert_images.py 04_workspace/[폴더명]/report_draft_[agent].docx
    python 03_code/insert_images.py input.docx output.docx --images-dir path/to/images

이미지 폴더는 기본적으로 입력 DOCX가 속한 작업 폴더의 images/ 하위를 사용한다.

삽입 로직:
  - 「※그림N은 별도 이미지로 삽입」 등의 플레이스홀더 단락을 이미지로 치환
  - 플레이스홀더가 없으면 「그림N」 캡션 단락의 직후에 삽입
  - 그림 번호 매핑: 1=causal.png, 2=kpi_tree.png (레거시 호환),
    그 외 번호는 images/ 내 나머지 PNG를 이름순으로 대응

종료코드: 0=전부 삽입, 1=실행 실패(입력/이미지 없음), 2=일부 또는 전무 삽입
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from _common import (
    EXIT_ERROR,
    EXIT_OK,
    EXIT_PARTIAL,
    WORKSPACE_DIR,
    configure_stdout,
    extract_target_id_from_docx,
)

IMAGE_WIDTH_MM = 145

# 「※그림N ... 이미지 ... 삽입」 / 「※図N ... 画像 ... 挿入」 플레이스홀더
PLACEHOLDER_RE = re.compile(r"※.*(?:図|그림)\s*(\d+).*(?:画像|이미지).*(?:挿入|삽입)")
# 「그림N: ...」 / 「図N: ...」 캡션
CAPTION_RE = re.compile(r"^(?:図|그림)\s*(\d+)[：:]")

# 레거시 그림 번호 → 파일명 매핑
LEGACY_IMAGE_NAMES = {1: "causal.png", 2: "kpi_tree.png"}


def find_figure_numbers(doc) -> set:
    """문서에서 참조되는 그림 번호를 수집"""
    numbers = set()
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = PLACEHOLDER_RE.search(text) or CAPTION_RE.search(text)
        if m:
            numbers.add(int(m.group(1)))
    return numbers


def build_image_map(img_dir: Path, figure_numbers: set) -> dict:
    """그림 번호 → 이미지 경로 매핑을 구성"""
    mapping = {}
    used_names = set()

    for n, name in LEGACY_IMAGE_NAMES.items():
        path = img_dir / name
        if path.exists():
            mapping[n] = path
            used_names.add(name)

    extra = [p for p in sorted(img_dir.glob("*.png")) if p.name not in used_names]
    unmapped = sorted(n for n in figure_numbers if n not in mapping)
    for n, path in zip(unmapped, extra):
        mapping[n] = path

    return mapping


def insert_image_at_paragraph(p, img_path, width_mm=IMAGE_WIDTH_MM):
    """단락의 텍스트를 지우고 이미지로 치환한다"""
    p.clear()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Mm(width_mm))


def add_image_after_paragraph(doc, p, img_path, caption_text, width_mm=IMAGE_WIDTH_MM):
    """단락 직후에 이미지와 캡션을 추가한다"""
    # 이미지 단락
    img_p = doc.add_paragraph()
    p._element.addnext(img_p._element)
    img_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = img_p.add_run()
    run.add_picture(str(img_path), width=Mm(width_mm))

    # 캡션 단락
    cap_p = doc.add_paragraph()
    img_p._element.addnext(cap_p._element)
    cap_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap_run = cap_p.add_run(caption_text)
    cap_run.font.size = Pt(9)


def insert_images(doc, image_map: dict) -> dict:
    """이미지를 삽입하고 번호별 성공 여부를 반환"""
    inserted = {n: False for n in image_map}

    # Phase 1: 플레이스홀더 단락을 이미지로 치환
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = PLACEHOLDER_RE.search(text)
        if m:
            n = int(m.group(1))
            if n in image_map and not inserted[n]:
                insert_image_at_paragraph(p, image_map[n])
                inserted[n] = True
                print(f"Inserted 그림{n} (placeholder replacement): {image_map[n].name}")

    # Phase 2: 플레이스홀더를 찾지 못한 번호는 캡션 행 직후에 삽입
    if not all(inserted.values()):
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            m = CAPTION_RE.search(text)
            if m:
                n = int(m.group(1))
                if n in image_map and not inserted[n]:
                    add_image_after_paragraph(doc, p, image_map[n], text)
                    inserted[n] = True
                    print(f"Inserted 그림{n} (after caption): {image_map[n].name}")

    return inserted


def main():
    configure_stdout()
    parser = argparse.ArgumentParser(description="DOCX 이미지 플레이스홀더 삽입")
    parser.add_argument("input", help="입력 DOCX 경로")
    parser.add_argument("output", nargs="?", help="출력 DOCX 경로 (기본: 입력 파일 덮어쓰기)")
    parser.add_argument("--images-dir", help="이미지 폴더 (기본: 작업 폴더의 images/)")
    args = parser.parse_args()

    docx_in = Path(args.input)
    if not docx_in.exists():
        print(f"ERROR: 입력 DOCX 없음: {docx_in}", file=sys.stderr)
        sys.exit(EXIT_ERROR)
    docx_out = Path(args.output) if args.output else docx_in

    if args.images_dir:
        img_dir = Path(args.images_dir)
    else:
        target_id = extract_target_id_from_docx(docx_in)
        print(f"분석 대상: {target_id}")
        img_dir = WORKSPACE_DIR / target_id / "images"

    if not img_dir.exists() or not any(img_dir.glob("*.png")):
        print(f"ERROR: 삽입할 이미지가 없음: {img_dir}", file=sys.stderr)
        sys.exit(EXIT_ERROR)

    doc = Document(str(docx_in))
    figure_numbers = find_figure_numbers(doc)
    image_map = build_image_map(img_dir, figure_numbers)

    if not image_map:
        print(f"ERROR: 그림 번호에 대응하는 이미지가 없음: {img_dir}", file=sys.stderr)
        sys.exit(EXIT_ERROR)

    inserted = insert_images(doc, image_map)

    failed = [n for n, done in inserted.items() if not done]
    for n in failed:
        print(f"WARNING: 그림{n} 미삽입 (문서에서 매칭 텍스트를 찾지 못함): {image_map[n].name}")

    doc.save(str(docx_out))
    print(f"Saved: {docx_out}")

    if failed:
        print(f"결과: [PARTIAL] {len(inserted) - len(failed)}/{len(inserted)}개 삽입")
        sys.exit(EXIT_PARTIAL)
    print(f"결과: [OK] {len(inserted)}개 전부 삽입")
    sys.exit(EXIT_OK)


if __name__ == "__main__":
    main()
