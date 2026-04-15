"""
Markdown to Word (.docx) Converter — 부동산 분석 보고서
STEP 12: Markdown→Word 변환 스크립트

사용법:
    python md_to_docx_converter.py [분석대상 ID]

예:
    python md_to_docx_converter.py seongsu-residential_KR
"""

import argparse
import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv

# ==================== 설정 ====================

# 포맷 설정
FONT_NAME = '맑은고딕'
FONT_SIZE_PT = 12
LINE_SPACING = 1.15
MARGIN_MM = 25
PARAGRAPH_SPACING_PT = 0

# 표지 설정(언어별 기본값)
COVER_PRESETS = {
    'ko': {
        'title': '부동산 분석 보고서',
        'date': '2026년 4월',
        'proposer_label': '작성자',
        'code_label': '분석 대상',
        'unknown_name': '대상',
        'unknown_loc': '미상',
        'input_filename': 'STEP11_보고서_draft.md',
        'output_filename': 'report_draft.docx',
        'master_csv': 'property_master_kr.csv',
    },
}
COVER_TITLE = COVER_PRESETS['ko']['title']
PROPOSAL_DATE = COVER_PRESETS['ko']['date']
PROPOSER_NAME = ''

# 경로 설정
ROOTDIR = Path(__file__).resolve().parent.parent
MIDDLE_OUTPUT_DIR = ROOTDIR / '04_workspace'
PROPERTY_MASTER_PATH = MIDDLE_OUTPUT_DIR / 'property_master.csv'

KNOWN_AGENT_TAGS = {'claude', 'codex'}

# ==================== 유틸리티 함수 ====================

def log(message, level='INFO'):
    """로그 출력"""
    try:
        print(f'[{level}] {message}')
    except UnicodeEncodeError:
        print(f'[{level}] {message.encode("ascii", "replace").decode()}')

def get_target_info(target_id, master_csv=None, lang='ko'):
    """분석대상 마스터에서 대상 정보를 가져옴(마스터 CSV가 없으면 target_id를 그대로 사용)"""
    preset = COVER_PRESETS.get(lang, COVER_PRESETS['ko'])
    unknown_name = preset['unknown_name']
    unknown_loc = preset['unknown_loc']

    # "_KR" 접미사를 제거한 순수 ID로 조회
    bare_id = re.sub(r'_KR$', '', str(target_id))

    candidates = []
    if master_csv:
        candidates.append(Path(master_csv) if os.path.isabs(master_csv) else MIDDLE_OUTPUT_DIR / master_csv)
    candidates.append(MIDDLE_OUTPUT_DIR / preset['master_csv'])
    candidates.append(PROPERTY_MASTER_PATH)

    for path in candidates:
        if not path.exists():
            continue
        try:
            with open(path, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    row_id = row.get('target_id') or row.get('property_id') or row.get('대상ID')
                    if row_id and str(row_id).strip() == bare_id:
                        return {
                            'code': bare_id,
                            'name': row.get('target_name') or row.get('property_name') or row.get('대상명') or f'{unknown_name}({bare_id})',
                            'location': row.get('location') or row.get('소재지') or unknown_loc,
                        }
        except Exception as e:
            log(f'분석대상 마스터 읽기 오류 ({path}): {e}', 'WARNING')

    return {'code': bare_id, 'name': bare_id, 'location': unknown_loc}

def count_chars(text):
    """문자수 카운트(공백·개행 제외)"""
    return len(re.sub(r'\s', '', text))


def extract_agent_tag(target_id):
    """target_id 끝의 에이전트 태그를 추출"""
    bare_id = re.sub(r'_KR$', '', str(target_id))
    parts = bare_id.split('_')
    if parts and parts[-1] in KNOWN_AGENT_TAGS:
        return parts[-1]
    return ''

# ==================== 포맷 설정 함수 ====================

def set_font(run, font_name=FONT_NAME, size_pt=FONT_SIZE_PT, bold=False, italic=False):
    """폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    # 동아시아 폰트 설정(중요)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_format(paragraph, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT,
                        line_spacing=LINE_SPACING, space_before=0, space_after=0):
    """단락 포맷 설정"""
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.first_line_indent = Pt(0)

def set_page_margins(section, margin_mm=MARGIN_MM):
    """페이지 여백 설정"""
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)

def set_cell_borders(cell):
    """표 셀 테두리 설정(검정·실선·0.5pt)"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5pt = 4 eighths of a point
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

    tcPr.append(tcBorders)

# ==================== Markdown 파싱 함수 ====================

def parse_inline_formatting(text):
    """인라인 서식(굵게, 이탤릭)을 파싱"""
    # 굵게: **text** or __text__
    # 이탤릭: *text* or _text_
    segments = []
    pattern = r'(\*\*|__)(.*?)\1|(\*|_)(.*?)\3'

    last_end = 0
    for match in re.finditer(pattern, text):
        # 매치 전의 텍스트
        if match.start() > last_end:
            segments.append({'text': text[last_end:match.start()], 'bold': False, 'italic': False})

        # 매치된 텍스트
        if match.group(1) in ['**', '__']:  # 굵게
            segments.append({'text': match.group(2), 'bold': True, 'italic': False})
        elif match.group(3) in ['*', '_']:  # 이탤릭
            segments.append({'text': match.group(4), 'bold': False, 'italic': True})

        last_end = match.end()

    # 남은 텍스트
    if last_end < len(text):
        segments.append({'text': text[last_end:], 'bold': False, 'italic': False})

    return segments if segments else [{'text': text, 'bold': False, 'italic': False}]

def parse_markdown_table(lines, start_idx):
    """Markdown 표를 파싱"""
    table_lines = []
    idx = start_idx

    # 표의 행을 수집
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        table_lines.append(line)
        idx += 1

    if len(table_lines) < 2:
        return None, start_idx

    # 헤더 행
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

    # 데이터 행(2번째 행은 구분자이므로 스킵)
    data = []
    for line in table_lines[2:]:
        row = [cell.strip() for cell in line.split('|')[1:-1]]
        if row:
            data.append(row)

    return {'header': header, 'data': data}, idx

def is_list_item(line):
    """불릿 항목 행인지 판정"""
    return re.match(r'^(\s*)([-*+])\s+(.+)$', line)

def get_list_level(line):
    """불릿 항목의 계층 레벨을 취득"""
    match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
    if match:
        indent = len(match.group(1))
        level = indent // 2  # 2스페이스당 1계층
        text = match.group(3)
        return level, text
    return 0, line

# ==================== Word 문서 생성 함수 ====================

def parse_cover_metadata(lines):
    """Markdown 앞부분의 표지 메타정보를 파싱하고, 스킵 후의 시작 행 index를 반환"""
    metadata = {}
    skip_until = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped == '---':
            skip_until = i + 1
            break
        m = re.match(r'\*\*분석 대상\*\*[:：]\s*(.+)', stripped)
        if m:
            metadata['target_name'] = m.group(1).strip()
        m = re.match(r'\*\*보고서 제목\*\*[:：]\s*(.+)', stripped)
        if m:
            metadata['proposal_title'] = m.group(1).strip()
    return metadata, skip_until


def create_cover_page(doc, target_info, cover_metadata=None, cover_config=None):
    """표지를 생성(왼쪽 정렬 비즈니스 리포트 스타일)"""
    log('표지 생성 중...')
    meta = cover_metadata or {}
    cfg = cover_config or {
        'title': COVER_TITLE,
        'date': PROPOSAL_DATE,
        'proposer': PROPOSER_NAME,
        'proposer_label': '작성자',
        'code_label': '분석 대상',
    }

    # 상단 여백용 빈 행
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_format(p)

    # 메인 타이틀
    p = doc.add_paragraph(cfg['title'])
    set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    set_font(p.runs[0], bold=True)

    # 서브 타이틀(제안서 타이틀)
    proposal_title = meta.get('proposal_title')
    if proposal_title:
        # em 대시로 분할되는 경우, 여러 행으로
        p = doc.add_paragraph(proposal_title)
        set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        set_font(p.runs[0])

    # 빈 행
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_format(p)

    # 분석 대상명
    name = meta.get('target_name', target_info['name'])
    p = doc.add_paragraph(name)
    set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    set_font(p.runs[0])

    # 분석 대상 ID
    p = doc.add_paragraph(f'{cfg["code_label"]}: {target_info["code"]}')
    set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    set_font(p.runs[0])

    # 하단 여백용 빈 행
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_format(p)

    # 제안일
    p = doc.add_paragraph(cfg['date'])
    set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
    set_font(p.runs[0])

    # 제안자
    proposer = (cfg.get('proposer') or '').strip()
    if proposer:
        p = doc.add_paragraph(f'{cfg["proposer_label"]}: {proposer}')
        set_paragraph_format(p, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        set_font(p.runs[0])

    # 페이지 나눔(표지 후에만)
    doc.add_page_break()
    log('표지 생성 완료')

def add_paragraph_with_formatting(doc, text):
    """서식 적용된 단락을 추가"""
    segments = parse_inline_formatting(text)
    p = doc.add_paragraph()

    for seg in segments:
        run = p.add_run(seg['text'])
        set_font(run, bold=seg['bold'], italic=seg['italic'])

    set_paragraph_format(p)
    return p

def add_table_from_markdown(doc, table_data):
    """Markdown 표를 Word 표로 변환"""
    header = table_data['header']
    data = table_data['data']

    rows = len(data) + 1  # 헤더 행 포함
    cols = len(header)

    table = doc.add_table(rows=rows, cols=cols)

    # 헤더 행
    for col_idx, header_text in enumerate(header):
        cell = table.rows[0].cells[col_idx]
        cell.text = ''
        paragraph = cell.paragraphs[0]
        # 헤더는 기본 굵게, 추가로 인라인 서식도 파싱
        segments = parse_inline_formatting(header_text)
        for seg in segments:
            run = paragraph.add_run(seg['text'])
            set_font(run, bold=True)
        set_paragraph_format(paragraph)
        set_cell_borders(cell)

    # 데이터 행
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < cols:
                cell = table.rows[row_idx].cells[col_idx]
                # 인라인 서식(굵게·이탤릭)을 파싱하여 적용
                cell.text = ''
                paragraph = cell.paragraphs[0]
                segments = parse_inline_formatting(cell_text)
                for seg in segments:
                    run = paragraph.add_run(seg['text'])
                    set_font(run, bold=seg['bold'], italic=seg['italic'])
                set_paragraph_format(paragraph)
                set_cell_borders(cell)

    log(f'표 처리 완료: {rows}행{cols}열')

    # 표 뒤에 빈 행 추가
    doc.add_paragraph()

def convert_markdown_to_docx(md_file_path, output_file_path, target_id, lang='ko', cover_override=None, master_csv=None):
    """Markdown 파일을 Word 문서로 변환"""

    # 파일 존재 확인
    if not os.path.exists(md_file_path):
        log(f'파일을 찾을 수 없음: {md_file_path}', 'ERROR')
        return False

    log(f'Markdown 파일 읽는 중: {md_file_path}')

    # Markdown 파일 읽기
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    log(f'읽기 완료: {len(lines)}행')

    # 분석대상 정보 취득
    target_info = get_target_info(target_id, master_csv=master_csv, lang=lang)
    log(f'분석 대상: {target_info["name"]} ({target_info["code"]})')

    # 표지 설정(언어 프리셋 + CLI 오버라이드)
    preset = COVER_PRESETS.get(lang, COVER_PRESETS['ko'])
    cover_config = {
        'title': preset['title'],
        'date': preset['date'],
        'proposer': PROPOSER_NAME,
        'proposer_label': preset['proposer_label'],
        'code_label': preset['code_label'],
    }
    if cover_override:
        cover_config.update({k: v for k, v in cover_override.items() if v})

    # Word 문서 작성
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    set_page_margins(section)

    # 기본 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_PT)

    # 표지 메타정보를 파싱(대상명·보고서 타이틀을 추출하고, 표지 부분을 스킵)
    cover_metadata, skip_lines = parse_cover_metadata(lines)
    log(f'표지 메타정보: {cover_metadata}, 스킵 행 수: {skip_lines}')

    # 표지 생성(메타정보·설정을 전달)
    create_cover_page(doc, target_info, cover_metadata, cover_config=cover_config)

    # 본문 처리(표지 부분은 스킵)
    idx = skip_lines
    total_chars = 0

    while idx < len(lines):
        line = lines[idx].rstrip()

        # 빈 행
        if not line:
            idx += 1
            continue

        # 제목1(# )
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.style = doc.styles['Heading 1']
            for run in p.runs:
                set_font(run, bold=True)
            set_paragraph_format(p)
            log(f'제목1 처리: {text}')
            total_chars += count_chars(text)
            idx += 1

        # 제목2(## )
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.style = doc.styles['Heading 2']
            for run in p.runs:
                set_font(run, bold=True)
            set_paragraph_format(p)
            log(f'제목2 처리: {text}')
            total_chars += count_chars(text)
            idx += 1

        # 제목3(### )
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.style = doc.styles['Heading 3']
            for run in p.runs:
                set_font(run, bold=True)
            set_paragraph_format(p)
            log(f'제목3 처리: {text}')
            total_chars += count_chars(text)
            idx += 1

        # 표(|로 시작)
        elif line.startswith('|'):
            table_data, next_idx = parse_markdown_table(lines, idx)
            if table_data:
                add_table_from_markdown(doc, table_data)
                # 표 내의 문자수를 카운트
                for row in table_data['data']:
                    for cell in row:
                        total_chars += count_chars(cell)
            idx = next_idx

        # 불릿 항목
        elif is_list_item(line):
            list_items = []
            while idx < len(lines) and is_list_item(lines[idx].rstrip()):
                level, text = get_list_level(lines[idx].rstrip())
                list_items.append((level, text))
                idx += 1

            for level, text in list_items:
                p = doc.add_paragraph(style='List Bullet')
                segments = parse_inline_formatting(text)
                for seg in segments:
                    run = p.add_run(seg['text'])
                    set_font(run, bold=seg['bold'], italic=seg['italic'])
                set_paragraph_format(p)
                # 들여쓰기 설정
                p.paragraph_format.left_indent = Mm(5 * level)
                total_chars += count_chars(text)

            log(f'불렛 처리: {len(list_items)}항목')

        # 일반 단락
        else:
            add_paragraph_with_formatting(doc, line)
            total_chars += count_chars(line)
            idx += 1

    # 문서 저장
    log(f'Word 문서 저장 중: {output_file_path}')
    doc.save(output_file_path)
    log(f'저장 완료: {output_file_path}')
    log(f'총 문자수(개산): {total_chars:,}자')

    return True

# ==================== 메인 처리 ====================

def main():
    parser = argparse.ArgumentParser(
        description='Markdown→Word 변환 — 부동산 분석 보고서',
    )
    parser.add_argument('target_id', help='분석대상 ID (예: seongsu-residential_KR, songpa-helio-city_KR)')
    parser.add_argument('--lang', choices=['ko'], default='ko', help='출력 언어 (기본: ko)')
    parser.add_argument('--input', help='입력 Markdown 파일명(기본: STEP11_보고서_draft.md)')
    parser.add_argument('--output', help='출력 DOCX 파일 경로')
    parser.add_argument('--master', help='분석대상 마스터 CSV(절대 경로 또는 04_workspace 바로 아래 상대 경로)')
    parser.add_argument('--title', help='표지 타이틀 덮어쓰기')
    parser.add_argument('--date', help='작성일 덮어쓰기')
    parser.add_argument('--proposer', help='작성자 덮어쓰기(기본: 표기 안 함)')
    args = parser.parse_args()

    target_id = args.target_id
    lang = args.lang
    preset = COVER_PRESETS[lang]
    agent_tag = extract_agent_tag(target_id)

    # 파일 경로 설정
    target_dir = MIDDLE_OUTPUT_DIR / target_id
    input_name = args.input or preset['input_filename']
    md_file = target_dir / input_name
    default_output_name = preset['output_filename']
    if agent_tag:
        default_output_name = f'report_draft_{agent_tag}.docx'
    output_file = Path(args.output) if args.output else target_dir / default_output_name

    cover_override = {
        'title': args.title,
        'date': args.date,
        'proposer': args.proposer,
    }

    log(f'========== 변환 시작: {target_id} (lang={lang}) ==========')
    log(f'입력 파일: {md_file}')
    log(f'출력 파일: {output_file}')

    # 변환 실행
    success = convert_markdown_to_docx(
        str(md_file), str(output_file), target_id,
        lang=lang, cover_override=cover_override, master_csv=args.master,
    )

    if success:
        log('========== 변환 성공 ==========', 'INFO')
        log(f'변환 완료: {output_file}')
        log('다음 단계: improve_docx_design.py 또는 insert_images.py 실행')
    else:
        log('========== 변환 실패 ==========', 'ERROR')
        print(f'\n❌ 변환 실패')
        sys.exit(1)

if __name__ == '__main__':
    main()
