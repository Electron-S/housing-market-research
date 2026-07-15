"""03_code 공용 유틸리티 — 직접 실행하지 않는다.

경로 상수, 작업 폴더명 파싱, cp949 안전 출력, 문자수 카운트 단일 정의,
종료코드 규약, DOCX 서식 헬퍼를 제공한다.

종료코드 규약 (전 스크립트 공통):
    0 = 완전 성공 / 검증 합격
    1 = 실행 실패 (입력 없음, 파싱/저장 오류)
    2 = 부분 실패 / 검증 불합격 (문자수 범위 밖, 일부 항목 실패 등)
"""
import re
import sys
from pathlib import Path
from typing import NamedTuple

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ==================== 경로 ====================

ROOTDIR = Path(__file__).resolve().parent.parent
WORKSPACE_DIR = ROOTDIR / '04_workspace'
OUTPUT_DIR = ROOTDIR / '05_output'

# ==================== 종료코드 ====================

EXIT_OK = 0
EXIT_ERROR = 1
EXIT_PARTIAL = 2

# ==================== 에이전트 태그 / 폴더명 파싱 ====================

KNOWN_AGENT_TAGS = frozenset({'claude', 'codex'})
_AGENT_TAG_ALT = '|'.join(sorted(KNOWN_AGENT_TAGS))
# report_designed_claude.docx 등 파일 stem 끝의 산출물 접미사
_DOCX_SUFFIX_RE = re.compile(r'_(?:designed|draft)(?:_(?:%s))?$' % _AGENT_TAG_ALT)


class TargetParts(NamedTuple):
    raw: str        # 예: "hanam-regenheim-comm_claude_KR"
    bare_id: str    # _KR 제거: "hanam-regenheim-comm_claude"
    agent_tag: str  # "claude" | "codex" | ""
    has_kr: bool    # _KR 접미사 여부


def parse_target(name) -> TargetParts:
    """폴더명/target_id에서 _KR 접미사와 에이전트 태그를 파싱하는 단일 진입점.

    한계: 규약(`[target_id]_[agent]_KR`)을 따르지 않는 이름
    (예: 하이픈 연결 태그, 태그 뒤 추가 접미사)에서는 태그를 감지하지 못한다.
    """
    raw = str(name)
    has_kr = raw.endswith('_KR')
    bare_id = raw[:-3] if has_kr else raw
    parts = bare_id.split('_')
    agent_tag = parts[-1] if parts and parts[-1] in KNOWN_AGENT_TAGS else ''
    return TargetParts(raw, bare_id, agent_tag, has_kr)


def strip_kr_suffix(name) -> str:
    """_KR 접미사를 제거한 이름을 반환."""
    return parse_target(name).bare_id


def extract_agent_tag(name) -> str:
    """폴더명/target_id 끝의 에이전트 태그를 반환 (없으면 '')."""
    return parse_target(name).agent_tag


def extract_agent_tag_from_filename(filepath) -> str:
    """파일명 stem 끝의 에이전트 태그를 반환 (예: report_draft_claude.docx -> 'claude')."""
    match = re.search(r'_(%s)$' % _AGENT_TAG_ALT, Path(filepath).stem)
    return match.group(1) if match else ''


def extract_target_id_from_docx(docx_path) -> str:
    """DOCX 경로에서 작업 폴더명(target_id)을 추출.

    부모 폴더가 04_workspace 바로 아래이면 폴더명을 신뢰하고,
    그 외에는 파일 stem에서 _designed/_draft(_에이전트) 접미사를 제거해 사용한다.
    """
    path = Path(docx_path).resolve()
    try:
        if path.parent.parent == WORKSPACE_DIR.resolve():
            return path.parent.name
    except (OSError, ValueError):
        pass
    return _DOCX_SUFFIX_RE.sub('', path.stem)


# ==================== cp949 안전 출력 ====================

def configure_stdout():
    """Windows cp949 콘솔 대비 — stdout/stderr를 errors='replace'로 재설정.

    각 스크립트 main() 첫 줄에서 호출한다.
    """
    for stream in (sys.stdout, sys.stderr):
        reconfigure = getattr(stream, 'reconfigure', None)
        if reconfigure is not None:
            try:
                reconfigure(errors='replace')
            except (ValueError, OSError):
                pass


def safe_print(*args, **kwargs):
    """print 래퍼 — UnicodeEncodeError 시 backslashreplace로 폴백."""
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        encoding = getattr(kwargs.get('file') or sys.stdout, 'encoding', None) or 'ascii'
        fallback = [
            str(a).encode(encoding, 'backslashreplace').decode(encoding, 'replace')
            for a in args
        ]
        print(*fallback, **kwargs)


def log(message, level='INFO'):
    """로그 출력 (cp949 안전)."""
    safe_print(f'[{level}] {message}')


# ==================== 문자수 (단일 정의) ====================
# 정본: DOCX 실측(공백 포함 total)으로 8,000~12,000자 게이트를 판정한다.
# Markdown 단계 수치는 예비 추정이며, 정식 판정은 count_docx_chars.py 결과를 따른다.

CHAR_LIMITS = {
    'ko': (8000, 12000),
}


class CharCount(NamedTuple):
    total: int      # 공백 포함 (게이트 판정 기준)
    no_space: int   # 공백(스페이스·전각공백) 제외 참고치


def _acc_line(text: str) -> CharCount:
    return CharCount(
        total=len(text),
        no_space=len(text.replace(' ', '').replace('　', '')),
    )


def count_text_chars(text: str) -> CharCount:
    """일반 텍스트의 문자수. 개행은 세지 않는다(DOCX 단락 카운트와 동일 축)."""
    total = 0
    no_space = 0
    for line in text.split('\n'):
        c = _acc_line(line)
        total += c.total
        no_space += c.no_space
    return CharCount(total, no_space)


def count_docx_chars(docx_path) -> CharCount:
    """DOCX 문자수 — 본문 + 표 + 머리글/바닥글 전부 포함 (게이트 정본 정의)."""
    from docx import Document

    doc = Document(str(docx_path))
    total = 0
    no_space = 0

    def acc(text):
        nonlocal total, no_space
        c = _acc_line(text)
        total += c.total
        no_space += c.no_space

    for p in doc.paragraphs:
        acc(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    acc(p.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            acc(p.text)
        for p in section.footer.paragraphs:
            acc(p.text)

    return CharCount(total, no_space)


def strip_markdown(text: str, include_tables: bool = True) -> str:
    """Markdown 문법 기호를 제거한 텍스트를 반환 (문자수 예비 추정용).

    - '#', '*' 기호와 수평선('---') 행 제거
    - include_tables=False이면 표 행('|' 시작)도 제외
    """
    lines = []
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped == '---':
            continue
        if not include_tables and stripped.startswith('|'):
            continue
        lines.append(re.sub(r'[#*]', '', line))
    return '\n'.join(lines)


# ==================== DOCX 서식 헬퍼 ====================

def set_cell_borders(cell, color='000000', size='4', clear_existing=False):
    """표 셀 4방향 테두리 설정 (실선, size는 1/8pt 단위: '4' = 0.5pt)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    if clear_existing:
        for border in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(border)

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_east_asia_font(obj, font_name):
    """run 또는 style의 동아시아 폰트(w:eastAsia)를 설정.

    font.name이 미설정이면 함께 설정한다 (rFonts 요소 생성 보장).
    """
    if obj.font.name is None:
        obj.font.name = font_name
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_run_font(run, font_name, size_pt=None, bold=None, italic=None, color=None):
    """run의 폰트명(+eastAsia)·크기·굵기·이탤릭·색을 한 번에 설정."""
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
