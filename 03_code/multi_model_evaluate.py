"""
STEP12 리뷰 패킷 생성 스크립트 — 부동산 분석 보고서

외부 API를 호출하지 않는다.
현재 실행 중인 에이전트(Claude Code / Codex)가 직접 보고서를 평가할 수 있도록
리뷰 패킷과 STEP12 결과 파일의 골격을 준비한다.

사용법:
    python 03_code/multi_model_evaluate.py <분석대상 ID>
    python 03_code/multi_model_evaluate.py seongsu-residential_KR
    python 03_code/multi_model_evaluate.py seongsu-residential_KR --reviewer codex
    python 03_code/multi_model_evaluate.py seongsu-residential_KR --iteration 2

입력 파일 탐색 순서:
    1. 04_workspace/{대상ID}/report_designed.docx
    2. 04_workspace/{대상ID}/report_draft.docx
    3. 04_workspace/{대상ID}/STEP11_보고서_draft.md
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

ROOTDIR = Path(__file__).resolve().parent.parent

CRITERIA = [
    ("데이터 신뢰성", "모든 주요 수치에 출처(국토부, KB, R-ONE 등)와 기준 시점이 명시됐는가"),
    ("시장 해석", "가격 흐름과 수요·공급 변화가 인과관계로 설명됐는가"),
    ("입지 분석", "교통·학군·업무 접근성이 실수요층과 구체적으로 연결됐는가"),
    ("리스크 반영", "대출 규제·입주 물량·금리 변화 등 하방 리스크가 구조화됐는가"),
    ("결론 실효성", "매수/보유/매도/개발/보류 등 행동 방향이 근거와 함께 명확히 제시됐는가"),
]

REVIEW_PACKET_NAME = "STEP12_review_packet.md"
STEP12_OUTPUT_NAME = "STEP12_output.md"


def extract_bare_id(target_id: str) -> str:
    return re.sub(r"_KR$", "", target_id)


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    parts = paragraphs
    if table_lines:
        parts += [""] + table_lines
    return "\n".join(parts)


def find_report_text(target_id: str) -> tuple[str, Path]:
    docx_candidates = [
        ROOTDIR / "04_workspace" / target_id / "report_designed.docx",
        ROOTDIR / "04_workspace" / target_id / "report_draft.docx",
        ROOTDIR / "05_output" / f"{extract_bare_id(target_id)}_designed.docx",
        ROOTDIR / "05_output" / f"{extract_bare_id(target_id)}.docx",
        ROOTDIR / "05_output" / f"{target_id}_designed.docx",
        ROOTDIR / "05_output" / f"{target_id}.docx",
    ]
    for path in docx_candidates:
        if path.exists():
            return extract_docx_text(path), path

    md_candidates = [
        ROOTDIR / "04_workspace" / target_id / "STEP11_보고서_draft.md",
        ROOTDIR / "04_workspace" / target_id / "STEP10_사업지분석_draft.md",
        ROOTDIR / "04_workspace" / target_id / "STEP10_output.md",
        ROOTDIR / "04_workspace" / target_id / "STEP11_output.md",
    ]
    for path in md_candidates:
        if path.exists():
            return path.read_text(encoding="utf-8"), path

    print(f"ERROR: 보고서를 찾을 수 없음: {target_id}")
    print("검색 경로:")
    for path in docx_candidates + md_candidates:
        print(f"  {path}")
    sys.exit(1)


def detect_next_iteration(output_path: Path) -> int:
    if not output_path.exists():
        return 1

    text = output_path.read_text(encoding="utf-8")
    matches = [int(value) for value in re.findall(r"^### Iteration (\d+)", text, flags=re.MULTILINE)]
    return (max(matches) + 1) if matches else 1


def ensure_step12_output_file(output_path: Path):
    if output_path.exists():
        return

    content = (
        "# STEP12 품질 리뷰\n\n"
        "이 파일은 Claude Code 또는 Codex가 직접 수행한 품질 리뷰 결과를 반복 단위로 누적 기록한다.\n"
        "각 반복은 `### Iteration N (YYYY-MM-DD)` 형식으로 추가한다.\n"
    )
    output_path.write_text(content, encoding="utf-8")


def build_summary_table_template() -> str:
    lines = [
        "| 평가항목 | 등급 | 코멘트 |",
        "|---|---|---|",
    ]
    for criterion, _ in CRITERIA:
        lines.append(f"| {criterion} | A/B/C | 코멘트 입력 |")
    return "\n".join(lines)


def build_output_template(iteration: int, today: str, reviewer: str, source_path: Path) -> str:
    return f"""### Iteration {iteration} ({today})

#### 메타정보
- Reviewer: {reviewer}
- Source: `{source_path}`

#### 종합평가
- Overall: A/B/C
- Reason: 한두 문장으로 요약

#### 평가결과 요약
{build_summary_table_template()}

#### 강점
- 
- 

#### 개선 필요 항목
- 
- 

#### 수정 지시
- 
- 

#### 반복 판정
- Status: PASS / REVISE
- Next Action: 다음 수정 또는 종료 기준
"""


def build_review_packet(
    target_id: str,
    reviewer: str,
    iteration: int,
    today: str,
    source_path: Path,
    report_text: str,
    output_path: Path,
) -> str:
    criteria_lines = "\n".join(
        f"| {idx} | {criterion} | {description} |"
        for idx, (criterion, description) in enumerate(CRITERIA, start=1)
    )

    return f"""# STEP12 Review Packet

## 메타정보
- Target ID: `{target_id}`
- Reviewer: `{reviewer}`
- Iteration: `{iteration}`
- Date: `{today}`
- Report Source: `{source_path}`
- Output File: `{output_path}`

## 목적
현재 실행 중인 에이전트가 보고서를 직접 읽고 5개 기준으로 평가한다.
외부 API를 호출하지 않는다.

## 평가 기준
| 번호 | 평가항목 | 고평가 기준 |
|---|---|---|
{criteria_lines}

## 리뷰 수행 지침
1. 보고서 전체를 읽고 먼저 종합 판단을 내린다.
2. 아래 5개 항목을 각각 A/B/C로 평가한다.
3. 점수가 낮은 항목은 왜 낮은지 구체적으로 적는다.
4. 수정해야 할 문장·데이터·논리 연결을 실무적으로 지시한다.
5. 모든 항목이 A가 아니면 `Status: REVISE`로 기록한다.
6. 수정 후에는 이 스크립트를 다시 실행해 다음 Iteration으로 반복한다.

## STEP12_output.md에 기록할 템플릿
아래 형식을 그대로 복사해 `{output_path.name}`에 추가한다.

~~~md
{build_output_template(iteration, today, reviewer, source_path)}
~~~

## 보고서 원문
아래 보고서를 기준으로 평가한다.

~~~text
{report_text}
~~~
"""


def write_review_packet(packet_path: Path, content: str):
    packet_path.write_text(content, encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="STEP12: 에이전트 기반 품질 리뷰 패킷 생성")
    parser.add_argument("target_id", help="분석대상 ID (예: seongsu-residential_KR)")
    parser.add_argument("--reviewer", default="codex", help="리뷰 수행 에이전트 이름 (예: codex, claude)")
    parser.add_argument("--iteration", type=int, help="강제 iteration 번호 (기본: STEP12_output.md 기준 자동 증가)")
    args = parser.parse_args()

    target_dir = ROOTDIR / "04_workspace" / args.target_id
    target_dir.mkdir(parents=True, exist_ok=True)

    step12_output_path = target_dir / STEP12_OUTPUT_NAME
    ensure_step12_output_file(step12_output_path)

    report_text, source_path = find_report_text(args.target_id)
    today = datetime.now().strftime("%Y-%m-%d")
    iteration = args.iteration if args.iteration is not None else detect_next_iteration(step12_output_path)

    packet_content = build_review_packet(
        target_id=args.target_id,
        reviewer=args.reviewer,
        iteration=iteration,
        today=today,
        source_path=source_path,
        report_text=report_text,
        output_path=step12_output_path,
    )

    packet_path = target_dir / REVIEW_PACKET_NAME
    write_review_packet(packet_path, packet_content)

    print(f"=== STEP12 리뷰 패킷 생성: {args.target_id} ===")
    print(f"보고서 소스: {source_path}")
    print(f"리뷰어: {args.reviewer}")
    print(f"Iteration: {iteration}")
    print(f"리뷰 패킷: {packet_path}")
    print(f"결과 기록 파일: {step12_output_path}")
    print("")
    print("다음 단계:")
    print(f"1. {packet_path.name}를 현재 에이전트에게 읽힌다.")
    print(f"2. 평가 결과를 {step12_output_path.name}에 Iteration {iteration}로 기록한다.")
    print("3. PASS가 아니면 보고서를 수정한 뒤 스크립트를 다시 실행한다.")


if __name__ == "__main__":
    main()
