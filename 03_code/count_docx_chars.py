"""DOCX 문자수 카운터 (크로스플랫폼, python-docx 기반)

- Windows 전용 count_docx_chars_win.py(Word COM) 의 macOS/Linux 호환 대체판
- 본문 + 표 + 머리글/바닥글 모두 카운트
- 언어별 규정 임계값 체크 (ko: 8,000~12,000)
- 종료코드: 0=OK, 1=범위 밖, 2=파일오류

사용 예:
    python 03_code/count_docx_chars.py 04_workspace/seongsu-residential_KR/report_designed.docx
    python 03_code/count_docx_chars.py path.docx --lang ko
    python 03_code/count_docx_chars.py path.docx --min 9000 --max 13000
"""
import argparse
import sys
from pathlib import Path

from docx import Document

# 언어별 규정 (파일럿 단계의 잠정치. 향후 실적에 따라 재조정 가능)
LIMITS = {
    'ko': (8000, 12000),  # 한글은 정보량 동일 시 다른언어 대비 약 60~80% 문자수
}


def count_chars(docx_path: Path):
    doc = Document(str(docx_path))
    total = 0
    total_no_space = 0

    def acc(text: str):
        nonlocal total, total_no_space
        total += len(text)
        total_no_space += len(text.replace(' ', '').replace('\u3000', ''))

    for p in doc.paragraphs:
        acc(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    acc(p.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            acc(p.text)
        for p in section.footer.paragraphs:
            acc(p.text)

    return total, total_no_space


def infer_lang(path: Path, cli_lang):
    if cli_lang:
        return cli_lang
    parts = path.resolve().parts
    name = str(path)
    if '_KR' in name or any('04_workspace' in p or '05_output' in p or p.endswith('_KR') for p in parts):
        return 'ko'
    return 'ko'


def main():
    parser = argparse.ArgumentParser(description='DOCX 문자수 카운터 (ko)')
    parser.add_argument('path', help='DOCX 파일 경로')
    parser.add_argument('--lang', choices=['ko'], help='언어 (미지정시 경로로 추정)')
    parser.add_argument('--min', type=int, help='하한 임계값 (기본 언어별)')
    parser.add_argument('--max', type=int, help='상한 임계값 (기본 언어별)')
    args = parser.parse_args()

    path = Path(args.path)
    if not path.exists():
        print(f'ERROR: 파일 없음: {path}', file=sys.stderr)
        sys.exit(2)

    lang = infer_lang(path, args.lang)
    lo, hi = LIMITS[lang]
    if args.min is not None:
        lo = args.min
    if args.max is not None:
        hi = args.max

    total, total_no_space = count_chars(path)

    print(f'=== Character Count: {path.name} ===')
    print(f'언어 기준: {lang}')
    print(f'문자수(공백포함): {total:,}')
    print(f'문자수(공백제외): {total_no_space:,}')
    print(f'규정 범위: {lo:,}~{hi:,} (공백포함 기준)')

    if lo <= total <= hi:
        print('결과: [OK] 범위 내')
        sys.exit(0)
    elif total < lo:
        print(f'결과: [UNDER] {lo - total:,}자 부족')
        sys.exit(1)
    else:
        print(f'결과: [OVER] {total - hi:,}자 초과')
        sys.exit(1)


if __name__ == '__main__':
    main()
